"""Filesystem MCP Tool.

Reads uploaded text files and extracts content.
Supports TXT, Markdown, PDF, and DOCX formats.
"""

import io
import logging
from typing import Any

logger = logging.getLogger(__name__)


class FilesystemTool:
    """Reads files and extracts text content.

    Supports:
    - Plain text (.txt)
    - Markdown (.md)
    - PDF (.pdf)
    - Word documents (.docx)
    """

    def read_file(self, file_path: str | None = None, content: bytes | None = None, filename: str = "") -> dict[str, Any]:
        """Read a file and extract text content.

        Auto-detects format from filename extension.

        Args:
            file_path: Path to file on disk.
            content: Raw file bytes (for uploads).
            filename: Original filename (for format detection).

        Returns:
            Dict with extracted text, format, and metadata.
        """
        ext = self._get_extension(file_path or filename)

        try:
            if ext in (".txt", ".md", ".markdown"):
                return self._read_text(file_path, content, ext)
            elif ext == ".pdf":
                return self._read_pdf(file_path, content)
            elif ext in (".docx", ".doc"):
                return self._read_docx(file_path, content)
            else:
                # Default: try reading as text
                return self._read_text(file_path, content, ext)
        except Exception as e:
            logger.error(f"Filesystem read error for '{filename}': {e}")
            return {"error": str(e), "source": "filesystem_tool", "filename": filename}

    def _get_extension(self, path: str) -> str:
        """Extract lowercase file extension."""
        if "." in path:
            return "." + path.rsplit(".", 1)[-1].lower()
        return ""

    def _read_text(self, file_path: str | None, content: bytes | None, ext: str) -> dict[str, Any]:
        """Read plain text or markdown file."""
        if content is not None:
            text = content.decode("utf-8", errors="replace")
        elif file_path is not None:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        else:
            return {"error": "No file path or content provided"}

        return {
            "source": "filesystem_tool",
            "format": "markdown" if ext in (".md", ".markdown") else "text",
            "text": text,
            "char_count": len(text),
            "line_count": text.count("\n") + 1,
        }

    def _read_pdf(self, file_path: str | None, content: bytes | None) -> dict[str, Any]:
        """Read PDF file and extract text."""
        try:
            import pypdf
        except ImportError:
            # Fallback: return info about the file without extracting
            return {
                "source": "filesystem_tool",
                "format": "pdf",
                "text": "[PDF content extraction requires pypdf. Install with: pip install pypdf]",
                "char_count": 0,
                "page_count": 0,
                "note": "pypdf not installed",
            }

        try:
            if content is not None:
                reader = pypdf.PdfReader(io.BytesIO(content))
            elif file_path is not None:
                reader = pypdf.PdfReader(file_path)
            else:
                return {"error": "No file path or content provided"}

            pages_text = []
            for page in reader.pages:
                pages_text.append(page.extract_text() or "")

            full_text = "\n\n".join(pages_text)

            return {
                "source": "filesystem_tool",
                "format": "pdf",
                "text": full_text,
                "char_count": len(full_text),
                "page_count": len(reader.pages),
            }
        except Exception as e:
            return {"error": f"PDF extraction failed: {e}", "source": "filesystem_tool"}

    def _read_docx(self, file_path: str | None, content: bytes | None) -> dict[str, Any]:
        """Read DOCX file and extract text."""
        try:
            import docx
        except ImportError:
            return {
                "source": "filesystem_tool",
                "format": "docx",
                "text": "[DOCX content extraction requires python-docx. Install with: pip install python-docx]",
                "char_count": 0,
                "note": "python-docx not installed",
            }

        try:
            if content is not None:
                doc = docx.Document(io.BytesIO(content))
            elif file_path is not None:
                doc = docx.Document(file_path)
            else:
                return {"error": "No file path or content provided"}

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)

            return {
                "source": "filesystem_tool",
                "format": "docx",
                "text": full_text,
                "char_count": len(full_text),
                "paragraph_count": len(paragraphs),
            }
        except Exception as e:
            return {"error": f"DOCX extraction failed: {e}", "source": "filesystem_tool"}
